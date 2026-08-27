from flask import Flask, render_template, request, jsonify, send_file, url_for
from flask_wtf.csrf import CSRFProtect
import sqlite3, os, io, re, json, uuid
from datetime import datetime
import ipaddress
import socket
from urllib.parse import urlparse, urljoin
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from xml.sax.saxutils import escape as xml_escape

app = Flask(__name__)
csrf = CSRFProtect(app)

app.config.update(
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE="Lax",
    MAX_CONTENT_LENGTH=20 * 1024 * 1024,
)


@app.errorhandler(400)
def handle_bad_request(e):
    # Surface CSRF failures (stale token after a server restart, an old
    # tab left open, etc.) as JSON so the frontend can show a clear
    # message instead of a generic "failed" toast.
    description = getattr(e, 'description', '') or ''
    if 'CSRF' in description.upper():
        return jsonify(
            error=(
                'Your session token expired (usually after the server '
                'restarts). Please refresh the page and try again.'
            )
        ), 400
    return jsonify(error=description or 'Bad request'), 400

@app.after_request
def add_security_headers(response):
    response.headers["X-Content-Type-Options"] = "nosniff"
    response.headers["X-Frame-Options"] = "SAMEORIGIN"
    response.headers["Referrer-Policy"] = "strict-origin-when-cross-origin"
    response.headers["Permissions-Policy"] = "camera=(), microphone=(), geolocation=()"
    return response


BASE = os.path.dirname(os.path.abspath(__file__))
DB_PATH = os.path.join(BASE, 'squid_bomber.db')
UPLOAD_DIR = os.path.join(BASE, 'static', 'uploads')
os.makedirs(UPLOAD_DIR, exist_ok=True)

CONTENT_TYPES = {'text', 'web', 'pdf', 'docx'}
ALLOWED_UPLOAD_EXTENSIONS = {'.pdf', '.docx'}
MAX_EXTRACT_IMAGES = 6
MAX_UPLOAD_IMAGES = 4
MAX_IMAGE_DIMENSION = 1600
MAX_UPLOAD_BYTES = 20 * 1024 * 1024

app.config['SECRET_KEY'] = os.environ.get('FLASK_SECRET_KEY')

if not app.config['SECRET_KEY']:
    raise RuntimeError(
        'FLASK_SECRET_KEY is not set. '
        'Set it as an environment variable before starting the application.'
    )

APP_ENV = os.environ.get('APP_ENV', 'development').lower()
IS_PRODUCTION = APP_ENV == 'production'

app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'
app.config['SESSION_COOKIE_SECURE'] = IS_PRODUCTION


INK_MODES = {
    'concept': {'label': 'Important Concept', 'color': '#a855f7'},
    'definition': {'label': 'Definition', 'color': '#38bdf8'},
    'example': {'label': 'Example / Application', 'color': '#4ade80'},
    'revision': {'label': 'Revision / Exam', 'color': '#fde047'},
    'doubt': {'label': 'Doubt', 'color': '#fb7185'},
    'task': {'label': 'Task / Check Later', 'color': '#fb923c'},
}


SAMPLE_QUIZ = [
    (
        'Which data structure follows LIFO?',
        ['Queue', 'Stack', 'Tree', 'Graph'],
        1
    ),
    (
        'Which language is used by Flask?',
        ['Python', 'Java', 'C++', 'Ruby'],
        0
    ),
    (
        'What does SQL primarily manage?',
        ['Images', 'Databases', 'CSS', 'Animations'],
        1
    ),
]


def is_safe_url(url):
    try:
        parsed = urlparse(url)

        if parsed.scheme not in ("http", "https"):
            return False

        if not parsed.hostname:
            return False

        hostname = parsed.hostname

        if hostname.lower() in (
            "localhost",
            "localhost.localdomain"
        ):
            return False

        addresses = socket.getaddrinfo(hostname, None)

        for address in addresses:
            ip = ipaddress.ip_address(address[4][0])

            if (
                ip.is_private
                or ip.is_loopback
                or ip.is_link_local
                or ip.is_reserved
                or ip.is_multicast
            ):
                return False

        return True

    except (ValueError, socket.gaierror, socket.herror):
        return False


def save_extracted_image(image_bytes, dest_dir, index):
    """Sanitize and persist an extracted image as PNG.

    Re-encoding through Pillow strips any malformed/malicious payload
    riding along inside the original image bytes and normalizes the
    format. Returns the saved filename, or None if the image could
    not be safely processed (never raises).
    """

    try:
        from PIL import Image

        img = Image.open(io.BytesIO(image_bytes))
        img.load()

        if img.width < 60 or img.height < 60:
            return None

        if img.mode != 'RGBA':
            img = img.convert('RGBA')

        if img.width > MAX_IMAGE_DIMENSION or img.height > MAX_IMAGE_DIMENSION:
            img.thumbnail((MAX_IMAGE_DIMENSION, MAX_IMAGE_DIMENSION))

        os.makedirs(dest_dir, exist_ok=True)

        filename = f'img{index}.png'

        img.save(os.path.join(dest_dir, filename), 'PNG')

        return filename

    except Exception:
        return None


def db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def init_db():
    conn = db()

    conn.executescript('''
    CREATE TABLE IF NOT EXISTS projects (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT NOT NULL,
        description TEXT DEFAULT '',
        created_at TEXT NOT NULL
    );

    CREATE TABLE IF NOT EXISTS notes (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        project_id INTEGER,
        title TEXT NOT NULL,
        body TEXT NOT NULL,
        created_at TEXT NOT NULL,
        updated_at TEXT NOT NULL,
        FOREIGN KEY(project_id) REFERENCES projects(id) ON DELETE CASCADE
    );

    CREATE TABLE IF NOT EXISTS highlights (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        project_id INTEGER,
        text TEXT NOT NULL,
        mode TEXT NOT NULL,
        color TEXT NOT NULL,
        created_at TEXT NOT NULL,
        FOREIGN KEY(project_id) REFERENCES projects(id) ON DELETE CASCADE
    );

    CREATE TABLE IF NOT EXISTS content (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        project_id INTEGER,
        title TEXT DEFAULT 'Study Material',
        source_url TEXT DEFAULT '',
        body TEXT NOT NULL,
        content_type TEXT DEFAULT 'text',
        images_json TEXT DEFAULT '[]',
        created_at TEXT NOT NULL,
        FOREIGN KEY(project_id) REFERENCES projects(id) ON DELETE CASCADE
    );

    CREATE TABLE IF NOT EXISTS quiz_attempts (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        project_id INTEGER,
        score INTEGER NOT NULL,
        total INTEGER NOT NULL,
        created_at TEXT NOT NULL,
        FOREIGN KEY(project_id) REFERENCES projects(id) ON DELETE CASCADE
    );
    ''')

    # Migrate older databases created before content_type/images_json existed.
    existing_cols = {
        row[1] for row in conn.execute(
            'PRAGMA table_info(content)'
        ).fetchall()
    }

    if 'content_type' not in existing_cols:
        conn.execute(
            "ALTER TABLE content ADD COLUMN content_type TEXT DEFAULT 'text'"
        )

    if 'images_json' not in existing_cols:
        conn.execute(
            "ALTER TABLE content ADD COLUMN images_json TEXT DEFAULT '[]'"
        )

    if conn.execute(
        'SELECT COUNT(*) FROM projects'
    ).fetchone()[0] == 0:

        now = datetime.utcnow().isoformat(timespec='seconds')

        conn.execute(
            'INSERT INTO projects(name, description, created_at) '
            'VALUES(?,?,?)',
            (
                'My First Study Project',
                'Start capturing knowledge here.',
                now
            )
        )

    conn.commit()
    conn.close()


@app.route('/')
def home():
    return render_template(
        'index.html',
        ink_modes=INK_MODES
    )


@app.get('/api/projects')
def get_projects():
    conn = db()

    rows = conn.execute(
        'SELECT * FROM projects ORDER BY id DESC'
    ).fetchall()

    conn.close()

    return jsonify([dict(r) for r in rows])


@app.post('/api/projects')
def create_project():
    data = request.get_json(silent=True) or {}

    name = (data.get('name') or '').strip()
    description = (data.get('description') or '').strip()

    if not name:
        return jsonify(
            error='Project name is required'
        ), 400

    if len(name) > 200:
        return jsonify(
            error='Project name is too long'
        ), 400

    if len(description) > 5000:
        return jsonify(
            error='Project description is too long'
        ), 400

    now = datetime.utcnow().isoformat(
        timespec='seconds'
    )

    conn = db()

    cur = conn.execute(
        'INSERT INTO projects(name, description, created_at) '
        'VALUES(?,?,?)',
        (
            name,
            description,
            now
        )
    )

    conn.commit()

    pid = cur.lastrowid

    conn.close()

    return jsonify(
        id=pid,
        name=name
    ), 201


@app.delete('/api/projects/<int:pid>')
def delete_project(pid):
    conn = db()

    conn.execute(
        'PRAGMA foreign_keys=ON'
    )

    conn.execute(
        'DELETE FROM projects WHERE id=?',
        (pid,)
    )

    conn.commit()
    conn.close()

    return jsonify(ok=True)


@app.get('/api/notes')
def get_notes():
    pid = request.args.get(
        'project_id',
        type=int
    )

    conn = db()

    rows = conn.execute(
        'SELECT * FROM notes '
        'WHERE project_id=? OR ? IS NULL '
        'ORDER BY id DESC',
        (
            pid,
            pid
        )
    ).fetchall()

    conn.close()

    return jsonify(
        [dict(r) for r in rows]
    )


@app.post('/api/notes')
def create_note():
    data = request.get_json(
        silent=True
    ) or {}

    title = (
        data.get('title')
        or 'Untitled Note'
    ).strip()

    body = (
        data.get('body')
        or ''
    ).strip()

    pid = data.get('project_id')

    if not body:
        return jsonify(
            error='Note body is required'
        ), 400

    if len(title) > 200:
        return jsonify(
            error='Note title is too long'
        ), 400

    if len(body) > 100000:
        return jsonify(
            error='Note body is too long'
        ), 400

    now = datetime.utcnow().isoformat(
        timespec='seconds'
    )

    conn = db()

    cur = conn.execute(
        'INSERT INTO notes('
        'project_id,title,body,created_at,updated_at'
        ') VALUES(?,?,?,?,?)',
        (
            pid,
            title,
            body,
            now,
            now
        )
    )

    conn.commit()

    nid = cur.lastrowid

    conn.close()

    return jsonify(
        id=nid
    ), 201


@app.put('/api/notes/<int:nid>')
def update_note(nid):
    data = request.get_json(
        silent=True
    ) or {}

    title = (
        data.get('title')
        or 'Untitled Note'
    ).strip()

    body = (
        data.get('body')
        or ''
    ).strip()

    if len(title) > 200:
        return jsonify(
            error='Note title is too long'
        ), 400

    if len(body) > 100000:
        return jsonify(
            error='Note body is too long'
        ), 400

    now = datetime.utcnow().isoformat(
        timespec='seconds'
    )

    conn = db()

    conn.execute(
        'UPDATE notes '
        'SET title=?, body=?, updated_at=? '
        'WHERE id=?',
        (
            title,
            body,
            now,
            nid
        )
    )

    conn.commit()
    conn.close()

    return jsonify(ok=True)


@app.delete('/api/notes/<int:nid>')
def delete_note(nid):
    conn = db()

    conn.execute(
        'DELETE FROM notes WHERE id=?',
        (nid,)
    )

    conn.commit()
    conn.close()

    return jsonify(ok=True)


@app.get('/api/highlights')
def get_highlights():
    pid = request.args.get(
        'project_id',
        type=int
    )

    conn = db()

    rows = conn.execute(
        'SELECT * FROM highlights '
        'WHERE project_id=? OR ? IS NULL '
        'ORDER BY id DESC',
        (
            pid,
            pid
        )
    ).fetchall()

    conn.close()

    return jsonify(
        [dict(r) for r in rows]
    )


@app.post('/api/highlights')
def create_highlight():
    data = request.get_json(
        silent=True
    ) or {}

    text = (
        data.get('text')
        or ''
    ).strip()

    mode = data.get(
        'mode',
        'concept'
    )

    if not text:
        return jsonify(
            error='Select some text first'
        ), 400

    if len(text) > 50000:
        return jsonify(
            error='Highlight text is too long'
        ), 400

    if mode not in INK_MODES:
        mode = 'concept'

    now = datetime.utcnow().isoformat(
        timespec='seconds'
    )

    conn = db()

    cur = conn.execute(
        'INSERT INTO highlights('
        'project_id,text,mode,color,created_at'
        ') VALUES(?,?,?,?,?)',
        (
            data.get('project_id'),
            text,
            mode,
            INK_MODES[mode]['color'],
            now
        )
    )

    conn.commit()

    hid = cur.lastrowid

    conn.close()

    return jsonify(
        id=hid
    ), 201


@app.delete('/api/highlights/<int:hid>')
def delete_highlight(hid):
    conn = db()

    conn.execute(
        'DELETE FROM highlights WHERE id=?',
        (hid,)
    )

    conn.commit()
    conn.close()

    return jsonify(ok=True)


@app.post('/api/content')
def save_content():
    data = request.get_json(
        silent=True
    ) or {}

    body = (
        data.get('body')
        or ''
    ).strip()

    title = (
        data.get('title')
        or 'Study Material'
    ).strip()

    source_url = (
        data.get('url')
        or ''
    ).strip()

    content_type = (
        data.get('content_type')
        or 'text'
    ).strip().lower()

    if content_type not in CONTENT_TYPES:
        content_type = 'text'

    raw_images = data.get('images')
    images = []

    if isinstance(raw_images, list):
        for item in raw_images[:MAX_EXTRACT_IMAGES]:
            if isinstance(item, str):
                item = item.strip()
                if item and len(item) <= 2048:
                    images.append(item)

    if not body:
        return jsonify(
            error='Content is empty'
        ), 400

    if len(title) > 200:
        return jsonify(
            error='Content title is too long'
        ), 400

    if len(source_url) > 2048:
        return jsonify(
            error='Source URL is too long'
        ), 400

    if len(body) > 500000:
        return jsonify(
            error='Content is too long'
        ), 400

    now = datetime.utcnow().isoformat(
        timespec='seconds'
    )

    conn = db()

    cur = conn.execute(
        'INSERT INTO content('
        'project_id,title,source_url,body,content_type,images_json,created_at'
        ') VALUES(?,?,?,?,?,?,?)',
        (
            data.get('project_id'),
            title,
            source_url,
            body,
            content_type,
            json.dumps(images),
            now
        )
    )

    conn.commit()

    cid = cur.lastrowid

    conn.close()

    return jsonify(
        id=cid
    ), 201


@app.get('/api/content')
def get_content():
    pid = request.args.get(
        'project_id',
        type=int
    )

    conn = db()

    rows = conn.execute(
        'SELECT * FROM content '
        'WHERE project_id=? OR ? IS NULL '
        'ORDER BY id DESC',
        (
            pid,
            pid
        )
    ).fetchall()

    conn.close()

    out = []

    for r in rows:
        d = dict(r)

        try:
            d['images'] = json.loads(
                d.pop('images_json', '[]') or '[]'
            )
        except Exception:
            d['images'] = []

        out.append(d)

    return jsonify(out)


@app.delete('/api/content/<int:cid>')
def delete_content(cid):
    conn = db()

    conn.execute(
        'DELETE FROM content WHERE id=?',
        (cid,)
    )

    conn.commit()
    conn.close()

    return jsonify(ok=True)


@app.post('/api/extract')
def extract():
    data = request.get_json(
        silent=True
    ) or {}

    url = (
        data.get('url')
        or ''
    ).strip()

    if not re.match(
        r'^https?://',
        url
    ):
        return jsonify(
            error='Enter a valid http/https URL'
        ), 400

    if not is_safe_url(url):
        return jsonify(
            error='This URL is not allowed'
        ), 400

    try:
        import requests
        from bs4 import BeautifulSoup

        try:
            from readability import Document
        except Exception:
            Document = None

        r = requests.get(
            url,
            timeout=12,
            headers={
                'User-Agent':
                    'SquidBomberStudyTool/1.0'
            }
        )

        r.raise_for_status()

        html = r.text

        if Document:
            doc = Document(html)

            title = (
                doc.short_title()
                or 'Study Material'
            )

            main_html = doc.summary()

            soup = BeautifulSoup(
                main_html,
                'html.parser'
            )

        else:
            soup = BeautifulSoup(
                html,
                'html.parser'
            )

            title = (
                soup.title.get_text(
                    ' ',
                    strip=True
                )
                if soup.title
                else 'Study Material'
            )

        # Collect a representative set of images from the page so
        # they can be shown alongside the extracted text. We never
        # download or proxy the bytes server-side — only absolute
        # http(s) URLs are returned, and the browser loads them
        # directly, so this adds no SSRF surface on the backend.
        def resolve_img_url(src):
            if not src:
                return None

            src = src.strip()

            if not src or src.startswith('data:'):
                return None

            absolute = urljoin(url, src)
            parsed = urlparse(absolute)

            if parsed.scheme not in ('http', 'https'):
                return None

            return absolute

        images = []
        seen = set()

        head_soup = BeautifulSoup(html, 'html.parser')

        meta_img = (
            head_soup.find('meta', attrs={'property': 'og:image'})
            or head_soup.find(
                'meta',
                attrs={'property': 'og:image:secure_url'}
            )
            or head_soup.find('meta', attrs={'name': 'twitter:image'})
        )

        primary = (
            resolve_img_url(meta_img.get('content'))
            if meta_img else None
        )

        if primary:
            images.append(primary)
            seen.add(primary)

        for img_tag in soup.find_all('img'):

            if len(images) >= MAX_EXTRACT_IMAGES:
                break

            src = (
                img_tag.get('src')
                or img_tag.get('data-src')
                or img_tag.get('data-lazy-src')
            )

            resolved = resolve_img_url(src)

            if not resolved or resolved in seen:
                continue

            try:
                w = int(img_tag.get('width') or 0)
                h = int(img_tag.get('height') or 0)

                if (w and w < 80) or (h and h < 80):
                    continue

            except ValueError:
                pass

            images.append(resolved)
            seen.add(resolved)

        images = images[:MAX_EXTRACT_IMAGES]

        for tag in soup(
            [
                'script',
                'style',
                'nav',
                'footer',
                'header',
                'aside',
                'form'
            ]
        ):
            tag.decompose()

        text = '\n\n'.join(
            x.strip()
            for x in soup.stripped_strings
            if len(x.strip()) > 2
        )

        text = re.sub(
            r'\n{3,}',
            '\n\n',
            text
        )

        return jsonify(
            title=title[:180],
            body=text[:50000],
            url=url,
            image=(images[0] if images else ''),
            images=images
        )

    except Exception as e:
        return jsonify(
            error=f'Could not extract this page: {e}'
        ), 502


@app.post('/api/extract/file')
def extract_file():
    f = request.files.get('file')

    if not f or not f.filename:
        return jsonify(
            error='Choose a PDF or DOCX file first'
        ), 400

    filename = f.filename
    ext = os.path.splitext(filename)[1].lower()

    if ext not in ALLOWED_UPLOAD_EXTENSIONS:
        return jsonify(
            error='Only .pdf and .docx files are supported'
        ), 400

    file_bytes = f.read()

    if not file_bytes:
        return jsonify(
            error='The uploaded file is empty'
        ), 400

    if len(file_bytes) > MAX_UPLOAD_BYTES:
        return jsonify(
            error='File is too large (max 20MB)'
        ), 400

    base_title = (
        os.path.splitext(os.path.basename(filename))[0][:180]
        or 'Uploaded Document'
    )

    subdir = uuid.uuid4().hex
    dest_dir = os.path.join(UPLOAD_DIR, subdir)
    images = []

    if ext == '.pdf':

        if not file_bytes.startswith(b'%PDF'):
            return jsonify(
                error='This does not look like a valid PDF file'
            ), 400

        try:
            from pypdf import PdfReader
        except Exception:
            return jsonify(
                error=(
                    'pypdf is not installed. '
                    'Run: pip install -r requirements.txt'
                )
            ), 500

        try:
            reader = PdfReader(io.BytesIO(file_bytes))

            if reader.is_encrypted:
                try:
                    reader.decrypt('')
                except Exception:
                    pass

            if reader.is_encrypted:
                return jsonify(
                    error='This PDF is password-protected'
                ), 400

            text_parts = []

            for page in reader.pages[:200]:
                try:
                    text_parts.append(
                        (page.extract_text() or '').strip()
                    )
                except Exception:
                    continue

            body = '\n\n'.join(t for t in text_parts if t)

            for page in reader.pages[:20]:

                if len(images) >= MAX_UPLOAD_IMAGES:
                    break

                try:
                    page_images = page.images
                except Exception:
                    continue

                for img in page_images:

                    if len(images) >= MAX_UPLOAD_IMAGES:
                        break

                    fn = save_extracted_image(
                        img.data,
                        dest_dir,
                        len(images) + 1
                    )

                    if fn:
                        images.append(
                            url_for(
                                'static',
                                filename=f'uploads/{subdir}/{fn}'
                            )
                        )

        except Exception as e:
            return jsonify(
                error=f'Could not read this PDF: {e}'
            ), 400

        content_type = 'pdf'

    else:

        if not file_bytes.startswith(b'PK\x03\x04'):
            return jsonify(
                error='This does not look like a valid DOCX file'
            ), 400

        try:
            docx_doc = Document(io.BytesIO(file_bytes))

            paragraphs = [
                p.text.strip()
                for p in docx_doc.paragraphs
                if p.text.strip()
            ]

            body = '\n\n'.join(paragraphs)

            for rel in docx_doc.part.rels.values():

                if len(images) >= MAX_UPLOAD_IMAGES:
                    break

                if 'image' not in rel.reltype:
                    continue

                try:
                    data = rel.target_part.blob
                except Exception:
                    continue

                fn = save_extracted_image(
                    data,
                    dest_dir,
                    len(images) + 1
                )

                if fn:
                    images.append(
                        url_for(
                            'static',
                            filename=f'uploads/{subdir}/{fn}'
                        )
                    )

        except Exception as e:
            return jsonify(
                error=f'Could not read this DOCX file: {e}'
            ), 400

        content_type = 'docx'

    if not body:
        body = '(No extractable text was found in this file.)'

    body = body[:200000]

    return jsonify(
        title=base_title,
        body=body,
        content_type=content_type,
        images=images,
        filename=filename
    )


@app.post('/api/quiz/<int:project_id>/submit')
def quiz_submit(project_id):
    data = request.get_json(
        silent=True
    ) or {}

    answers = data.get(
        'answers',
        []
    )

    if not isinstance(
        answers,
        list
    ):
        return jsonify(
            error='Answers must be a list'
        ), 400

    if len(answers) > len(SAMPLE_QUIZ):
        return jsonify(
            error='Too many answers submitted'
        ), 400

    score = 0

    for i, ans in enumerate(answers):

        try:
            answer_index = int(ans)

        except (
            TypeError,
            ValueError
        ):
            return jsonify(
                error='Invalid quiz answer'
            ), 400

        if (
            i < len(SAMPLE_QUIZ)
            and answer_index
            == SAMPLE_QUIZ[i][2]
        ):
            score += 1

    conn = db()

    now = datetime.utcnow().isoformat(
        timespec='seconds'
    )

    conn.execute(
        'INSERT INTO quiz_attempts('
        'project_id,score,total,created_at'
        ') VALUES(?,?,?,?)',
        (
            project_id,
            score,
            len(SAMPLE_QUIZ),
            now
        )
    )

    conn.commit()
    conn.close()

    return jsonify(
        score=score,
        total=len(SAMPLE_QUIZ)
    )


@app.get('/api/quiz')
def quiz():
    return jsonify(
        [
            {
                'question': q,
                'options': opts
            }
            for q, opts, _
            in SAMPLE_QUIZ
        ]
    )


@app.get('/api/analytics')
def analytics():
    conn = db()

    out = {}

    for key, table in [
        ('projects', 'projects'),
        ('notes', 'notes'),
        ('highlights', 'highlights')
    ]:
        out[key] = conn.execute(
            f'SELECT COUNT(*) FROM {table}'
        ).fetchone()[0]

    out['quiz_attempts'] = conn.execute(
        'SELECT COUNT(*) FROM quiz_attempts'
    ).fetchone()[0]

    row = conn.execute(
        'SELECT COALESCE('
        'AVG(score*100.0/NULLIF(total,0)),'
        '0'
        ') FROM quiz_attempts'
    ).fetchone()

    out['avg_score'] = round(
        row[0],
        1
    )

    conn.close()

    return jsonify(out)


@app.post('/api/translate')
def translate_text():

    data = request.get_json(
        silent=True
    ) or {}

    text = (
        data.get('text')
        or ''
    ).strip()

    target = (
        data.get('target_lang')
        or ''
    ).strip()

    if not text:
        return jsonify(
            error='No text to translate'
        ), 400

    if not target:
        return jsonify(
            error='Target language is required'
        ), 400

    if len(text) > 50000:
        return jsonify(
            error='Text is too long to translate'
        ), 400

    try:
        from deep_translator import GoogleTranslator

    except Exception:
        return jsonify(
            error=(
                'deep-translator is not installed. '
                'Run: pip install -r requirements.txt'
            )
        ), 500

    try:
        # deep-translator's free backend has a per-request
        # character limit, so split on paragraph boundaries
        # and translate in safe-sized chunks.
        chunks = []
        chunk = ''

        for para in text.split('\n'):

            if len(chunk) + len(para) + 1 > 4500:
                chunks.append(chunk)
                chunk = para
            else:
                chunk = (chunk + '\n' + para) if chunk else para

        if chunk:
            chunks.append(chunk)

        translator = GoogleTranslator(
            source='auto',
            target=target
        )

        translated_parts = [
            translator.translate(c) or ''
            for c in chunks
        ]

        translated = '\n'.join(translated_parts)

        return jsonify(
            translated=translated,
            target_lang=target
        )

    except Exception as e:
        return jsonify(
            error=f'Translation failed: {e}'
        ), 502


@app.post('/api/export/pdf')
def export_pdf():

    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import (
            SimpleDocTemplate,
            Paragraph,
            Spacer
        )
        from reportlab.lib.styles import (
            getSampleStyleSheet
        )
        from reportlab.lib.enums import (
            TA_CENTER
        )

    except Exception:
        return jsonify(
            error=(
                'ReportLab is not installed. '
                'Run: pip install -r requirements.txt'
            )
        ), 500

    data = request.get_json(
        silent=True
    ) or {}

    pid = data.get(
        'project_id'
    )

    conn = db()

    project = conn.execute(
        'SELECT * FROM projects WHERE id=?',
        (pid,)
    ).fetchone()

    notes = conn.execute(
        'SELECT * FROM notes '
        'WHERE project_id=? '
        'ORDER BY id DESC',
        (pid,)
    ).fetchall()

    highlights = conn.execute(
        'SELECT * FROM highlights '
        'WHERE project_id=? '
        'ORDER BY id DESC',
        (pid,)
    ).fetchall()

    conn.close()

    buf = io.BytesIO()

    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        rightMargin=40,
        leftMargin=40,
        topMargin=40,
        bottomMargin=40
    )

    styles = getSampleStyleSheet()

    styles['Title'].alignment = TA_CENTER

    story = []

    title = (
        project['name']
        if project
        else 'Squid Bomber Study Material'
    )

    story.append(
        Paragraph(
            title,
            styles['Title']
        )
    )

    story.append(
        Spacer(
            1,
            20
        )
    )

    story.append(
        Paragraph(
            'Highlights',
            styles['Heading2']
        )
    )

    for x in highlights:

        label = INK_MODES.get(
            x['mode'],
            {}
        ).get(
            'label',
            x['mode']
        )

        story.append(
            Paragraph(
                f"<b>{xml_escape(label)}</b>: "
                f"{xml_escape(x['text']).replace(chr(10), '<br/>')}",
                styles['BodyText']
            )
        )

        story.append(
            Spacer(
                1,
                8
            )
        )

    story.append(
        Paragraph(
            'Notes',
            styles['Heading2']
        )
    )

    for n in notes:

        story.append(
            Paragraph(
                f"<b>{xml_escape(n['title'])}</b><br/>"
                f"{xml_escape(n['body']).replace(chr(10), '<br/>')}",
                styles['BodyText']
            )
        )

        story.append(
            Spacer(
                1,
                10
            )
        )

    doc.build(story)

    buf.seek(0)

    return send_file(
        buf,
        as_attachment=True,
        download_name='squid_bomber_export.pdf',
        mimetype='application/pdf'
    )


@app.context_processor
def inject():
    return {
        'ink_modes': INK_MODES
    }


init_db()


if __name__ == '__main__':
    debug_mode = (
        os.environ.get(
            'FLASK_DEBUG',
            '0'
        ) == '1'
    )

    app.run(
        debug=debug_mode,
        host='127.0.0.1',
        port=5000
    )